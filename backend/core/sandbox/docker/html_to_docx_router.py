#!/usr/bin/env python3
import json
import os
from pathlib import Path
from typing import Dict, Optional
import tempfile
import re
from io import BytesIO

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError as e:
    raise ImportError(f"python-docx is not installed. Please install it with: pip install python-docx. Error: {e}")

from bs4 import BeautifulSoup


router = APIRouter(prefix="/document", tags=["docx-conversion"])

output_dir = Path("generated_docx")
output_dir.mkdir(exist_ok=True)


class ConvertRequest(BaseModel):
    doc_path: str = Field(..., description="Path to the document file (.doc for TipTap documents)")
    download: bool = Field(False, description="If true, returns the DOCX file directly. If false, returns JSON with download URL.")


class ConvertResponse(BaseModel):
    success: bool
    message: str
    docx_url: Optional[str] = None
    filename: Optional[str] = None


class HTMLToDocxConverter:
    def __init__(self, doc_path: str):
        self.doc_path = Path(doc_path)
        self.doc_data = {}
        self.document = None
        
    def load_document(self):
        if not self.doc_path.exists():
            raise FileNotFoundError(f"Document not found: {self.doc_path}")
            
        with open(self.doc_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        try:
            self.doc_data = json.loads(content)
            if self.doc_data.get('type') != 'tiptap_document':
                raise ValueError("Not a valid TipTap document")
        except json.JSONDecodeError:
            self.doc_data = {
                'title': self.doc_path.stem,
                'content': content,
                'metadata': {}
            }
    
    def create_docx(self) -> Document:
        self.document = Document()

        html_content = self.doc_data.get('content', '')
        self.convert_html_to_docx(html_content)
        
        return self.document
    
    def convert_html_to_docx(self, html_content: str):
        if not html_content:
            return
            
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for element in soup.children:
            self.process_element(element)
    
    def process_element(self, element, parent_paragraph=None):
        if element.name is None:
            text = str(element).strip()
            if text and parent_paragraph:
                parent_paragraph.add_run(text)
            return

        if element.name == 'p':
            p = self.document.add_paragraph()
            for child in element.children:
                self.process_inline_element(child, p)
                
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            heading = self.document.add_heading(element.get_text(), level)
            
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = self.document.add_paragraph(style='List Bullet')
                for child in li.children:
                    self.process_inline_element(child, p)
                    
        elif element.name == 'ol':
            for i, li in enumerate(element.find_all('li', recursive=False), 1):
                p = self.document.add_paragraph(style='List Number')
                for child in li.children:
                    self.process_inline_element(child, p)
                    
        elif element.name == 'blockquote':
            p = self.document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            run = p.add_run(element.get_text())
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            
        elif element.name == 'pre':
            code_text = element.get_text()
            p = self.document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(64, 64, 64)
            
        elif element.name == 'hr':
            p = self.document.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif element.name == 'br':
            if parent_paragraph:
                parent_paragraph.add_run('\n')
            else:
                self.document.add_paragraph()
                
        elif element.name == 'table':
            self.process_table(element)
            
        elif element.name == 'img':
            # Image - add as text reference for now
            alt_text = element.get('alt', 'Image')
            src = element.get('src', '')
            p = self.document.add_paragraph()
            p.add_run(f'[Image: {alt_text}]').italic = True
            if src:
                p.add_run(f' ({src})').font.size = Pt(8)
    
    def process_inline_element(self, element, paragraph):
        if element.name is None:
            text = str(element).strip()
            if text:
                paragraph.add_run(text)
            return
        
        if element.name == 'strong' or element.name == 'b':
            run = paragraph.add_run(element.get_text())
            run.bold = True
            
        elif element.name == 'em' or element.name == 'i':
            run = paragraph.add_run(element.get_text())
            run.italic = True
            
        elif element.name == 'u':
            run = paragraph.add_run(element.get_text())
            run.underline = True
            
        elif element.name == 's':
            run = paragraph.add_run(element.get_text())
            run.font.strike = True
            
        elif element.name == 'code':
            run = paragraph.add_run(element.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 0, 128)
            
        elif element.name == 'a':
            text = element.get_text()
            href = element.get('href', '')
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
            if href:
                paragraph.add_run(f' ({href})').font.size = Pt(8)
                
        elif element.name == 'br':
            paragraph.add_run('\n')
            
        else:
            for child in element.children:
                self.process_inline_element(child, paragraph)
    
    def process_table(self, table_element):
        rows = table_element.find_all('tr')
        if not rows:
            return
            
        max_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cells))
        
        if max_cols == 0:
            return
        
        
        table = self.document.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
       
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    table_cell = table.rows[i].cells[j]
                    table_cell.text = cell.get_text().strip()
                    
                    if cell.name == 'th':
                        for paragraph in table_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    async def convert_to_docx(self, store_locally: bool = True) -> tuple:
        self.load_document()
        
        doc = self.create_docx()
        
        doc_title = self.doc_data.get('title', 'document')
        safe_title = re.sub(r'[^\w\s-]', '', doc_title.lower())
        safe_title = re.sub(r'[-\s]+', '-', safe_title)[:50]
        
        if store_locally:
            docx_path = output_dir / f"{safe_title}.docx"
            doc.save(str(docx_path))
            return docx_path, safe_title
        else:
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read(), safe_title


@router.post("/convert-to-docx")
async def convert_document_to_docx(request: ConvertRequest):
    """
    Convert TipTap HTML document to DOCX.
    
    Takes a document file path and returns either:
    - DOCX file directly (if download=true) - uses document title as filename
    - JSON response with download URL (if download=false, default)
    """
    try:
        # Validate document path exists
        doc_path = Path(request.doc_path)
        if not doc_path.exists():
            raise HTTPException(status_code=404, detail=f"Document not found: {request.doc_path}")
        
        # Create converter
        converter = HTMLToDocxConverter(request.doc_path)
        
        # If download is requested, return file directly
        if request.download:
            docx_content, doc_name = await converter.convert_to_docx(store_locally=False)
            
            return Response(
                content=docx_content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=\"{doc_name}.docx\""}
            )
        
        # Otherwise, store locally and return JSON with download URL
        docx_path, doc_name = await converter.convert_to_docx(store_locally=True)
        
        docx_url = f"/downloads/{docx_path.name}"
        
        return ConvertResponse(
            success=True,
            message=f"DOCX generated successfully",
            docx_url=docx_url,
            filename=docx_path.name
        )
        
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"❌ Conversion error: {e}")
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")


@router.get("/health")
async def docx_health_check():
    """DOCX service health check endpoint."""
    return {"status": "healthy", "service": "docx-converter"}
